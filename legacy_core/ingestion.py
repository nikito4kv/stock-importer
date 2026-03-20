from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document

from .common import normalize_whitespace


@dataclass(slots=True)
class IngestionIssue:
    code: str
    level: str
    message: str
    original_index: int | None = None
    paragraph_no: int | None = None


@dataclass(slots=True)
class ParagraphRecord:
    paragraph_no: int
    original_index: int
    text: str
    numbering_valid: bool
    numbering_style: str

    def as_payload(self) -> dict[str, str | int | bool]:
        return {
            "paragraph_no": self.paragraph_no,
            "original_index": self.original_index,
            "text": self.text,
            "numbering_valid": self.numbering_valid,
        }


@dataclass(slots=True)
class ScriptIngestionResult:
    source_file: Path
    header_text: str
    paragraphs: list[ParagraphRecord]
    issues: list[IngestionIssue]


_EXPLICIT_NUMBER_RE = re.compile(r"^(?P<number>\d+)\s*[\.)]\s*(?P<body>.*)$")
_HEADING_PREFIXES = (
    "part ",
    "chapter ",
    "section ",
    "scene ",
    "act ",
    "episode ",
)
_HEADING_TITLES = {
    "introduction",
    "prologue",
    "epilogue",
    "conclusion",
    "outro",
    "finale",
}


def _paragraph_has_word_numbering(paragraph: Any) -> bool:
    element = getattr(paragraph, "_element", None)
    ppr = getattr(element, "pPr", None)
    return bool(ppr is not None and ppr.numPr is not None)


def _paragraph_style_name(paragraph: Any) -> str:
    style = getattr(paragraph, "style", None)
    name = getattr(style, "name", "")
    return str(name or "")


def _looks_like_heading(text: str, paragraph: Any) -> bool:
    normalized = normalize_whitespace(text)
    if not normalized:
        return False

    lowered = normalized.casefold()
    if lowered in _HEADING_TITLES or any(
        lowered.startswith(prefix) for prefix in _HEADING_PREFIXES
    ):
        return True

    style_name = _paragraph_style_name(paragraph).casefold()
    if style_name.startswith("heading") or style_name in {"title", "subtitle"}:
        return True

    words = normalized.split()
    letters = [char for char in normalized if char.isalpha()]
    uppercase_letters = [char for char in letters if char.isupper()]
    mostly_uppercase = bool(letters) and len(uppercase_letters) / len(letters) >= 0.8
    title_like = (
        len(words) <= 12
        and len(normalized) <= 120
        and not re.search(r"[.!?]$", normalized)
    )
    return title_like and mostly_uppercase


def normalize_paragraph_payload(
    *,
    paragraph_no: int,
    original_index: int,
    text: str,
    numbering_valid: bool,
) -> dict[str, str | int | bool]:
    return {
        "paragraph_no": int(paragraph_no),
        "original_index": int(original_index),
        "text": normalize_whitespace(text),
        "numbering_valid": bool(numbering_valid),
    }


def ingest_script_docx(file_path: str | Path) -> ScriptIngestionResult:
    path = Path(file_path)
    if not path.exists() or not path.is_file():
        raise FileNotFoundError(f"File not found: {path}")

    doc = Document(str(path))
    header_parts: list[str] = []
    paragraphs: list[ParagraphRecord] = []
    issues: list[IngestionIssue] = []
    started_content = False
    expected_number = 1

    for index, para in enumerate(doc.paragraphs, start=1):
        text = normalize_whitespace(para.text)
        if not text:
            continue

        match = _EXPLICIT_NUMBER_RE.match(text)
        explicit_number = int(match.group("number")) if match else None
        body_text = normalize_whitespace(match.group("body")) if match else text
        has_word_numbering = _paragraph_has_word_numbering(para)
        is_numbered = explicit_number is not None or has_word_numbering
        is_heading = _looks_like_heading(text, para)

        if not is_numbered:
            if is_heading:
                if not started_content:
                    header_parts.append(text)
                continue

            started_content = True
            paragraphs.append(
                ParagraphRecord(
                    paragraph_no=expected_number,
                    original_index=index,
                    text=text,
                    numbering_valid=True,
                    numbering_style="implicit",
                )
            )
            expected_number += 1
            continue

        started_content = True
        paragraph_no = (
            explicit_number if explicit_number is not None else expected_number
        )
        numbering_valid = True
        numbering_style = "explicit" if explicit_number is not None else "word"

        if paragraph_no != expected_number:
            numbering_valid = False
            issues.append(
                IngestionIssue(
                    code="invalid_numbering_sequence",
                    level="error",
                    message=(
                        f"Expected paragraph number {expected_number}, got {paragraph_no}."
                    ),
                    original_index=index,
                    paragraph_no=paragraph_no,
                )
            )

        if paragraph_no == 1 and expected_number != 1:
            numbering_valid = False

        if expected_number == 1 and paragraph_no != 1:
            numbering_valid = False
            issues.append(
                IngestionIssue(
                    code="numbering_must_start_at_one",
                    level="error",
                    message="Numbered script paragraphs must start at 1.",
                    original_index=index,
                    paragraph_no=paragraph_no,
                )
            )

        if explicit_number is not None and not body_text:
            numbering_valid = False
            issues.append(
                IngestionIssue(
                    code="empty_numbered_paragraph",
                    level="error",
                    message="Numbered paragraph content cannot be empty.",
                    original_index=index,
                    paragraph_no=paragraph_no,
                )
            )

        paragraphs.append(
            ParagraphRecord(
                paragraph_no=paragraph_no,
                original_index=index,
                text=text,
                numbering_valid=numbering_valid,
                numbering_style=numbering_style,
            )
        )
        expected_number = paragraph_no + 1

    if not paragraphs:
        issues.append(
            IngestionIssue(
                code="no_script_paragraphs",
                level="error",
                message="No usable script paragraphs were found in the document.",
            )
        )

    return ScriptIngestionResult(
        source_file=path,
        header_text="\n".join(header_parts).strip(),
        paragraphs=paragraphs,
        issues=issues,
    )


def read_script_paragraphs(
    file_path: str | Path,
) -> tuple[str, list[dict[str, str | int | bool]]]:
    result = ingest_script_docx(file_path)
    payload = [
        normalize_paragraph_payload(
            paragraph_no=paragraph.paragraph_no,
            original_index=paragraph.original_index,
            text=paragraph.text,
            numbering_valid=paragraph.numbering_valid,
        )
        for paragraph in result.paragraphs
    ]
    return result.header_text, payload
