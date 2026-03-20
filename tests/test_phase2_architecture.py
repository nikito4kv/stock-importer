from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document

from app.bootstrap import bootstrap_application
from app.runtime import DesktopApplication
from config.settings import SUPPORTED_DESKTOP_STACK
from domain.enums import RunStatus, SessionHealth
from domain.models import AssetSelection, ParagraphUnit, Preset
from storage.repositories import SettingsRepository
from storage.serialization import read_json, write_json
from storage.workspace import WorkspaceStorage


def _write_docx(directory: str, paragraphs: list[str]) -> Path:
    path = Path(directory) / "script.docx"
    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(str(path))
    return path


class Phase2ArchitectureTests(unittest.TestCase):
    def test_desktop_application_bootstrap_and_project_creation(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            script_path = _write_docx(
                temp_dir, ["HEADER", "1. First scene", "2. Second scene"]
            )

            application = DesktopApplication.create(temp_dir)
            snapshot = application.start()

            self.assertEqual(snapshot.workspace_root, Path(temp_dir))
            self.assertIn("storyblocks_video", snapshot.providers)
            self.assertTrue((Path(temp_dir) / "projects").exists())
            self.assertTrue((Path(temp_dir) / "runs").exists())

            project = application.create_project("Demo", script_path)
            reloaded = application.container.project_repository.load(project.project_id)

            self.assertIsNotNone(reloaded)
            assert reloaded is not None
            assert reloaded.script_document is not None
            self.assertEqual(reloaded.name, "Demo")
            self.assertEqual(reloaded.script_document.header_text, "HEADER")
            self.assertEqual(
                [item.paragraph_no for item in reloaded.script_document.paragraphs],
                [1, 2],
            )
            self.assertTrue(reloaded.script_document.paragraphs[0].intent is not None)
            self.assertTrue(
                reloaded.script_document.paragraphs[0].query_bundle is not None
            )

    def test_settings_manager_applies_nested_preset_and_secret_round_trip(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            container = bootstrap_application(temp_dir)

            preset = Preset(
                name="fast-lane",
                settings_snapshot={
                    "desktop_stack": "tk",
                    "concurrency": {"paragraph_workers": 5, "queue_size": 3},
                    "browser": {"slow_mode": False, "action_delay_ms": 250},
                },
            )
            container.settings_manager.save_preset(preset)

            applied = container.settings_manager.apply_preset(
                container.settings, preset.name
            )
            self.assertEqual(applied.concurrency.paragraph_workers, 5)
            self.assertEqual(applied.concurrency.queue_size, 3)
            self.assertFalse(applied.browser.slow_mode)
            self.assertEqual(applied.browser.action_delay_ms, 250)
            self.assertEqual(applied.desktop_stack, SUPPORTED_DESKTOP_STACK)

            container.settings_manager.set_secret("gemini_test", "secret-value")
            self.assertEqual(
                container.settings_manager.get_secret("gemini_test"), "secret-value"
            )

    def test_settings_repository_normalizes_legacy_desktop_stack(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            paths = WorkspaceStorage(temp_dir).initialize()
            repository = SettingsRepository(paths)
            settings_path = paths.config_dir / "settings.json"
            write_json(settings_path, {"desktop_stack": "tk"})

            settings = repository.load()
            repository.save(settings)
            persisted = read_json(settings_path)

            self.assertEqual(settings.desktop_stack, SUPPORTED_DESKTOP_STACK)
            self.assertEqual(persisted.get("desktop_stack"), SUPPORTED_DESKTOP_STACK)

    def test_desktop_launch_uses_qt_only_path(self) -> None:
        import ui

        launched: list[object] = []
        controller = object()
        with patch(
            "ui.import_module",
            return_value=SimpleNamespace(launch_pyside_app=launched.append),
        ) as import_module_mock:
            ui.launch_desktop_app(controller)

        import_module_mock.assert_called_once_with("ui.qt_app")
        self.assertEqual(launched, [controller])

    def test_desktop_launch_raises_clear_error_when_pyside6_missing(self) -> None:
        import ui

        missing_qt = ModuleNotFoundError("No module named 'PySide6'")
        missing_qt.name = "PySide6"

        with patch("ui.import_module", side_effect=missing_qt):
            with self.assertRaisesRegex(RuntimeError, "PySide6 is required"):
                ui.launch_desktop_app(object())

    def test_browser_profile_registry_lifecycle(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            container = bootstrap_application(temp_dir)

            profile = container.profile_registry.create_profile(
                "Primary",
                container.workspace.paths.browser_profiles_dir,
            )
            self.assertTrue(profile.storage_path.exists())

            renamed = container.profile_registry.rename_profile(
                profile.profile_id, "Storyblocks Main"
            )
            self.assertEqual(renamed.display_name, "Storyblocks Main")

            active = container.profile_registry.set_active(profile.profile_id)
            self.assertTrue(active.is_active)

            state = container.session_manager.set_health(
                profile.profile_id, SessionHealth.READY
            )
            self.assertEqual(state.health, SessionHealth.READY)
            self.assertEqual(
                container.session_manager.current_state().profile_id, profile.profile_id
            )

            profile_json = (
                container.workspace.paths.browser_profiles_dir
                / f"{profile.profile_id}.json"
            )
            profile_dir = profile.storage_path
            container.profile_registry.delete_profile(profile.profile_id)
            self.assertFalse(profile_json.exists())
            self.assertFalse(profile_dir.exists())

    def test_media_run_service_uses_configured_shared_orchestrator(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            container = bootstrap_application(temp_dir)

            self.assertIs(
                container.media_run_service._orchestrator, container.orchestrator
            )

    def test_orchestrator_pause_resume_and_failure_tracking(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            container = bootstrap_application(temp_dir)
            paragraphs = [
                ParagraphUnit(paragraph_no=1, original_index=1, text="One"),
                ParagraphUnit(paragraph_no=2, original_index=2, text="Two"),
                ParagraphUnit(paragraph_no=3, original_index=3, text="Three"),
            ]

            def processor(paragraph: ParagraphUnit) -> AssetSelection:
                return AssetSelection(
                    paragraph_no=paragraph.paragraph_no, status="done"
                )

            run = container.orchestrator.create_run("project-1")
            container.orchestrator.pause_after_current(run.run_id)
            paused = container.orchestrator.execute(run, paragraphs, processor)

            self.assertEqual(paused.status, RunStatus.PAUSED)
            self.assertEqual(len(paused.completed_paragraphs), 1)

            resumed = container.orchestrator.resume(
                paused.run_id, paragraphs, processor
            )
            self.assertEqual(resumed.status, RunStatus.COMPLETED)
            self.assertEqual(sorted(resumed.completed_paragraphs), [1, 2, 3])

            failed_run = container.orchestrator.create_run("project-2")

            def flaky_processor(paragraph: ParagraphUnit) -> AssetSelection:
                if paragraph.paragraph_no == 2:
                    raise RuntimeError("provider failed")
                return AssetSelection(
                    paragraph_no=paragraph.paragraph_no, status="done"
                )

            failed = container.orchestrator.execute(
                failed_run, paragraphs, flaky_processor
            )
            self.assertEqual(failed.status, RunStatus.FAILED)
            self.assertEqual(sorted(failed.completed_paragraphs), [1, 3])
            self.assertEqual(failed.failed_paragraphs, [2])
            self.assertEqual(failed.last_error, "provider failed")
            self.assertIn(
                "paragraph.failed",
                [event.name for event in container.event_recorder.events],
            )


if __name__ == "__main__":
    unittest.main()
