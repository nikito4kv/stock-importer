from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document

from app.runtime import DesktopApplication
from domain.models import ParagraphIntent, ParagraphUnit, QueryBundle, ScriptDocument
from pipeline.ingestion import ScriptIngestionService
from pipeline.intents import ParagraphIntentService


def _write_docx(directory: str, paragraphs: list[str]) -> Path:
    path = Path(directory) / "script.docx"
    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(str(path))
    return path


class FakeModel:
    def __init__(self, responses: list[str]):
        self._responses = list(responses)

    def generate_content(self, prompt: str) -> SimpleNamespace:
        if not self._responses:
            raise AssertionError(f"Unexpected extra prompt: {prompt}")
        return SimpleNamespace(text=self._responses.pop(0))


class ParagraphIntentTests(unittest.TestCase):
    @staticmethod
    def _word_count(value: str) -> int:
        return len([token for token in value.split() if token])

    def test_bootstrap_paragraph_intent_limits_queries_to_two_words(self) -> None:
        service = ParagraphIntentService()

        intent, query_bundle = service.bootstrap_paragraph_intent(
            paragraph_no=1,
            paragraph_text=(
                "Ancient explorers paddle a damaged canoe through a foggy jungle river at dawn."
            ),
            strictness="balanced",
        )

        self.assertTrue(intent.primary_video_queries)
        self.assertTrue(intent.image_queries)
        for query in intent.primary_video_queries + intent.image_queries:
            self.assertLessEqual(self._word_count(query), 2)
        for queries in query_bundle.provider_queries.values():
            for query in queries:
                self.assertLessEqual(self._word_count(query), 2)

    def test_parse_intent_response_replaces_abstract_video_queries(self) -> None:
        service = ParagraphIntentService()
        intent = service.parse_intent_response(
            json.dumps(
                {
                    "subject": "Spanish soldiers",
                    "action": "rowing",
                    "setting": "jungle river",
                    "mood": "fearful",
                    "style": "cinematic",
                    "negative_terms": ["logo", "illustration"],
                    "source_language": "en",
                    "translated_queries": [],
                    "estimated_duration_seconds": 12,
                    "primary_video_queries": ["fear of the unknown", "alien force"],
                    "image_queries": [
                        "submissive victims",
                        "Spanish soldiers on jungle river",
                    ],
                }
            ),
            paragraph_no=47,
            paragraph_text="Fear of the unknown paralyzed the exhausted Spanish soldiers rowing through a jungle river.",
            strictness="balanced",
        )

        query_bundle = service.build_query_bundle(
            intent,
            strictness="balanced",
        )

        self.assertEqual(intent.primary_video_queries[0], "Spanish soldiers")
        self.assertNotIn("fear of the unknown", intent.primary_video_queries)
        self.assertNotIn("alien force", intent.primary_video_queries)
        self.assertIn("jungle river", intent.image_queries)
        self.assertEqual(
            query_bundle.provider_queries["storyblocks_video"][0], "Spanish soldiers"
        )
        self.assertTrue(
            any(
                query.startswith("rowing") or query.endswith("photo")
                for query in query_bundle.provider_queries["free_image"]
            )
        )
        for queries in query_bundle.provider_queries.values():
            for query in queries:
                self.assertLessEqual(self._word_count(query), 2)

    def test_bad_low_value_gemini_phrases_are_replaced_with_concrete_visual_queries(
        self,
    ) -> None:
        service = ParagraphIntentService()
        intent = service.parse_intent_response(
            json.dumps(
                {
                    "subject": "armed warriors",
                    "action": "approaching canoes",
                    "setting": "stone temple",
                    "primary_video_queries": [
                        "what told",
                        "powerful civilization",
                        "situation required",
                    ],
                    "image_queries": ["what everything", "expedition did"],
                }
            ),
            paragraph_no=5,
            paragraph_text=(
                "Armed warriors watch wooden canoes approach a stone temple on the river bank at dawn."
            ),
            strictness="balanced",
        )

        query_bundle = service.build_query_bundle(intent, strictness="balanced")

        self.assertEqual(intent.primary_video_queries[0], "armed warriors")
        self.assertIn(
            "stone temple", query_bundle.provider_queries["storyblocks_image"]
        )
        disallowed = {
            "what told",
            "powerful civilization",
            "situation required",
            "what everything",
            "expedition did",
        }
        for queries in query_bundle.provider_queries.values():
            self.assertTrue(disallowed.isdisjoint(set(queries)))

    def test_single_generic_nouns_are_deprioritized_even_if_model_returns_them(
        self,
    ) -> None:
        service = ParagraphIntentService()
        intent = service.parse_intent_response(
            json.dumps(
                {
                    "subject": "ancient settlers",
                    "action": "rowing canoe",
                    "setting": "river bank",
                    "primary_video_queries": ["civilization", "expedition"],
                    "image_queries": ["civilization", "expedition"],
                }
            ),
            paragraph_no=8,
            paragraph_text=(
                "Ancient settlers row a wooden canoe toward the river bank near a stone shrine."
            ),
            strictness="balanced",
        )

        self.assertNotIn("civilization", intent.primary_video_queries)
        self.assertNotIn("expedition", intent.primary_video_queries)
        self.assertIn("river bank", intent.primary_video_queries)

    def test_build_prompt_forbids_opening_words_and_helper_phrases(self) -> None:
        prompt = ParagraphIntentService().build_prompt(
            3,
            "What told the expedition what everything meant near a powerful civilization.",
            strictness="balanced",
        )

        self.assertIn("Do not copy the opening words mechanically", prompt)
        self.assertIn("what told", prompt)
        self.assertIn("1 or 2 English words only", prompt)

    def test_extract_document_builds_golden_intents_for_multiple_paragraphs(
        self,
    ) -> None:
        service = ParagraphIntentService()
        with tempfile.TemporaryDirectory() as temp_dir:
            path = _write_docx(
                temp_dir,
                [
                    "INTRODUCTION",
                    "1. Spanish soldiers row a damaged boat through a misty jungle river.",
                    "2. Orellana writes in his diary while the expedition reaches the open sea.",
                ],
            )
            document = ScriptIngestionService().ingest(path)
            model = FakeModel(
                [
                    json.dumps(
                        {
                            "subject": "Spanish soldiers",
                            "action": "rowing damaged boat",
                            "setting": "misty jungle river",
                            "mood": "tense",
                            "style": "documentary",
                            "negative_terms": ["illustration"],
                            "source_language": "en",
                            "translated_queries": [],
                            "estimated_duration_seconds": 11,
                            "primary_video_queries": [
                                "Spanish soldiers rowing damaged boat on misty jungle river",
                                "fear of uncertainty",
                            ],
                            "image_queries": [
                                "Spanish soldiers damaged boat jungle river",
                                "misty jungle river boat",
                            ],
                        }
                    ),
                    json.dumps(
                        {
                            "subject": "Francisco de Orellana",
                            "action": "writing in diary",
                            "setting": "open sea boat deck",
                            "mood": "exhausted",
                            "style": "historical documentary",
                            "negative_terms": ["map", "text overlay"],
                            "source_language": "ru",
                            "translated_queries": [
                                "Francisco de Orellana writing diary on boat deck",
                                "expedition reaches open sea",
                            ],
                            "estimated_duration_seconds": 18,
                            "primary_video_queries": [
                                "Francisco de Orellana writing diary on boat deck",
                                "open sea expedition boat",
                            ],
                            "image_queries": [
                                "historical explorer writing diary on deck",
                                "open sea expedition boat",
                            ],
                        }
                    ),
                ]
            )

            intents_by_paragraph, items, updated_document = service.extract_document(
                model,
                document,
                strictness="strict",
                max_workers=1,
            )

            self.assertEqual(sorted(intents_by_paragraph.keys()), [1, 2])
            self.assertEqual(updated_document.header_text, "INTRODUCTION")
            self.assertEqual(len(items), 2)
            self.assertIsNotNone(updated_document.paragraphs[0].query_bundle)
            self.assertIsNotNone(updated_document.paragraphs[1].query_bundle)
            self.assertIsNotNone(updated_document.paragraphs[1].intent)
            assert updated_document.paragraphs[0].query_bundle is not None
            assert updated_document.paragraphs[1].query_bundle is not None
            assert updated_document.paragraphs[1].intent is not None
            self.assertLessEqual(
                len(updated_document.paragraphs[0].query_bundle.video_queries), 2
            )
            for paragraph in updated_document.paragraphs:
                assert paragraph.query_bundle is not None
                for queries in paragraph.query_bundle.provider_queries.values():
                    for query in queries:
                        self.assertLessEqual(self._word_count(query), 2)
            self.assertEqual(
                updated_document.paragraphs[1].intent.source_language, "ru"
            )
            self.assertEqual(
                updated_document.paragraphs[1].intent.estimated_duration_seconds, 18.0
            )
            self.assertEqual(
                set(updated_document.paragraphs[1].query_bundle.provider_queries),
                {"storyblocks_video", "storyblocks_image", "free_image"},
            )

    def test_extract_document_collects_intent_timing_aggregates_and_error_counters(
        self,
    ) -> None:
        service = ParagraphIntentService()
        document = ScriptDocument(
            source_path=Path("story.docx"),
            header_text="HEADER",
            paragraphs=[
                ParagraphUnit(
                    paragraph_no=1,
                    original_index=1,
                    text="Spanish soldiers row a damaged boat through a jungle river.",
                ),
                ParagraphUnit(
                    paragraph_no=2,
                    original_index=2,
                    text="The exhausted crew drifts in silence near the shoreline.",
                ),
            ],
        )
        model = FakeModel(
            [
                json.dumps(
                    {
                        "subject": "Spanish soldiers",
                        "action": "rowing",
                        "setting": "jungle river",
                        "primary_video_queries": ["Spanish soldiers", "jungle river"],
                        "image_queries": ["jungle river", "damaged boat"],
                    }
                ),
                "not json",
                "still not json",
                "nope",
            ]
        )

        _intents, items, _updated_document = service.extract_document(
            model,
            document,
            strictness="balanced",
            max_workers=1,
        )
        metrics = service.last_extract_metrics()

        self.assertEqual(len(items), 2)
        self.assertIn("metrics", items[0])
        self.assertIn("metrics", items[1])
        self.assertIn("intent_total_ms", items[0]["metrics"])
        self.assertIn("intent_total_ms", items[1]["metrics"])
        self.assertEqual(metrics["intent_errors_total"], 1)
        self.assertIn("intent_p50_ms", metrics)
        self.assertIn("intent_p95_ms", metrics)
        self.assertIn("intent_error_types", metrics)

    def test_extract_document_preserves_original_error_type_in_metrics(self) -> None:
        service = ParagraphIntentService()
        document = ScriptDocument(
            source_path=Path("story.docx"),
            header_text="HEADER",
            paragraphs=[
                ParagraphUnit(
                    paragraph_no=1,
                    original_index=1,
                    text="A crew rows through dense jungle fog.",
                )
            ],
        )
        model = FakeModel(
            [
                json.dumps({"subject": "crew", "action": "rowing", "setting": "river"}),
                json.dumps({"subject": "crew", "action": "rowing", "setting": "river"}),
                json.dumps({"subject": "crew", "action": "rowing", "setting": "river"}),
            ]
        )

        with patch.object(
            service,
            "parse_intent_response",
            side_effect=RuntimeError("synthetic parse failure"),
        ):
            _intents, _items, _updated_document = service.extract_document(
                model,
                document,
                strictness="balanced",
                max_workers=1,
            )

        metrics = service.last_extract_metrics()
        self.assertEqual(metrics["intent_errors_total"], 1)
        self.assertEqual(metrics["intent_error_types"].get("RuntimeError"), 1)

    def test_save_intents_json_writes_new_contract(self) -> None:
        service = ParagraphIntentService()
        document = ScriptDocument(
            source_path=Path("story.docx"),
            header_text="INTRO",
            paragraphs=[
                ParagraphUnit(
                    paragraph_no=1,
                    original_index=2,
                    text="A damaged boat drifts through the jungle river.",
                )
            ],
        )
        paragraph = service.apply_manual_edit(
            document.paragraphs[0],
            intent=ParagraphIntent(
                paragraph_no=1,
                subject="damaged boat",
                action="drifting",
                setting="jungle river",
                mood="tense",
                style="documentary",
                negative_terms=["illustration"],
                source_language="en",
                translated_queries=[],
                estimated_duration_seconds=9,
                primary_video_queries=["damaged boat drifting jungle river"],
                image_queries=["damaged boat on jungle river"],
            ),
            strictness="balanced",
        )
        items = [service.build_item_payload(paragraph)]

        with tempfile.TemporaryDirectory() as temp_dir:
            out_file = service.save_intents_json(
                document,
                items,
                output_path=Path(temp_dir) / "paragraph_intents.json",
                model_name="fake-gemini",
                strictness="balanced",
            )
            payload = json.loads(out_file.read_text(encoding="utf-8"))

        self.assertEqual(payload["contract"], "paragraph_intents")
        self.assertEqual(payload["schema_version"], 2)
        self.assertEqual(payload["header_text"], "INTRO")
        self.assertEqual(payload["items"][0]["paragraph_no"], 1)
        self.assertEqual(
            payload["items"][0]["query_bundle"]["provider_queries"][
                "storyblocks_video"
            ][0],
            "damaged boat",
        )
        self.assertTrue(
            all(not str(key).startswith("include_") for key in payload.keys())
        )

    def test_desktop_application_persists_manual_intent_edit(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            script_path = _write_docx(
                temp_dir,
                ["INTRO", "1. A damaged boat drifts through the jungle river."],
            )
            app = DesktopApplication.create(temp_dir)
            project = app.create_project("Phase 3", script_path)

            intent = ParagraphIntent(
                paragraph_no=1,
                subject="damaged boat",
                action="drifting",
                setting="jungle river",
                mood="tense",
                style="documentary",
                negative_terms=["illustration"],
                source_language="en",
                translated_queries=[],
                estimated_duration_seconds=9.0,
                primary_video_queries=["damaged boat drifting jungle river"],
                image_queries=["damaged boat on jungle river"],
            )

            updated = app.update_paragraph_intent(
                project.project_id,
                1,
                intent=intent,
                query_bundle=QueryBundle(
                    video_queries=["damaged boat drifting jungle river"],
                    image_queries=["damaged boat jungle river photo"],
                    provider_queries={
                        "storyblocks_video": ["damaged boat drifting jungle river"],
                        "storyblocks_image": ["damaged boat on jungle river"],
                        "free_image": ["damaged boat jungle river photo"],
                    },
                ),
            )

            reloaded = app.container.project_repository.load(updated.project_id)
            self.assertIsNotNone(reloaded)
            assert reloaded is not None
            assert reloaded.script_document is not None
            paragraph = reloaded.script_document.paragraphs[0]
            assert paragraph.intent is not None
            assert paragraph.query_bundle is not None
            self.assertEqual(paragraph.intent.subject, "damaged boat")
            self.assertEqual(
                paragraph.query_bundle.provider_queries["storyblocks_video"][0],
                "damaged boat drifting jungle river",
            )

    def test_build_prompt_includes_manual_prompt_and_full_script_context(self) -> None:
        service = ParagraphIntentService()

        prompt = service.build_prompt(
            3,
            "River boat at dawn.",
            strictness="balanced",
            manual_prompt="Prefer archival realism",
            full_script_context="Header: INTRO\n1. Jungle river\n2. Boat crew",
        )

        self.assertIn("Additional operator guidance", prompt)
        self.assertIn("Prefer archival realism", prompt)
        self.assertIn("Optional full script context", prompt)
        self.assertIn("Header: INTRO", prompt)

    def test_desktop_application_uses_gemini_only_for_explicit_enrichment(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            script_path = _write_docx(
                temp_dir,
                [
                    "INTRO",
                    "1. The survivors push a broken raft through a flooded jungle passage.",
                ],
            )
            app = DesktopApplication.create(temp_dir)
            app.container.settings_manager.set_secret(
                app.container.settings.security.gemini_api_key_secret_name,
                "test-gemini-key-1234567890",
            )
            model = FakeModel(
                [
                    json.dumps(
                        {
                            "subject": "survivor group",
                            "action": "pushing raft",
                            "setting": "jungle floodplain",
                            "mood": "tense",
                            "style": "documentary",
                            "negative_terms": ["illustration"],
                            "source_language": "en",
                            "translated_queries": [],
                            "estimated_duration_seconds": 10,
                            "primary_video_queries": [
                                "survivor group",
                                "broken raft",
                                "jungle floodplain",
                            ],
                            "image_queries": [
                                "broken raft",
                                "jungle floodplain",
                            ],
                        }
                    )
                ]
            )

            with patch(
                "app.runtime.create_gemini_model", return_value=model
            ) as factory:
                project = app.create_project("Gemini Project", script_path)
                factory.assert_not_called()
                enriched = app.enrich_project_intents(
                    project.project_id,
                    strictness="balanced",
                    manual_prompt="Prefer documentary footage",
                    attach_full_script_context=True,
                )

            factory.assert_called_once()
            assert enriched.script_document is not None
            paragraph = enriched.script_document.paragraphs[0]
            assert paragraph.query_bundle is not None
            self.assertEqual(
                paragraph.query_bundle.provider_queries["storyblocks_video"][0],
                "survivor group",
            )
            for queries in paragraph.query_bundle.provider_queries.values():
                for query in queries:
                    self.assertLessEqual(self._word_count(query), 2)


if __name__ == "__main__":
    unittest.main()
