from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

import keyword_extractor
from legacy_core.ingestion import ingest_script_docx


def _write_docx(directory: str, paragraphs: list[str]) -> Path:
    path = Path(directory) / "script.docx"
    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(str(path))
    return path


class IngestionTests(unittest.TestCase):
    def test_ingest_numbered_docx_with_header(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = _write_docx(
                temp_dir, ["INTRODUCTION", "1. First scene", "2) Second scene"]
            )
            result = ingest_script_docx(path)

            self.assertEqual(result.header_text, "INTRODUCTION")
            self.assertEqual(
                [paragraph.paragraph_no for paragraph in result.paragraphs], [1, 2]
            )
            self.assertEqual(
                [paragraph.original_index for paragraph in result.paragraphs], [2, 3]
            )
            self.assertEqual([issue.level for issue in result.issues], [])

    def test_ingest_reports_invalid_numbering(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = _write_docx(temp_dir, ["2. Wrong start", "4. Gap"])
            result = ingest_script_docx(path)

            codes = {issue.code for issue in result.issues}
            self.assertIn("numbering_must_start_at_one", codes)
            self.assertIn("invalid_numbering_sequence", codes)

    def test_ingest_skips_section_headings_between_paragraphs(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = _write_docx(
                temp_dir,
                [
                    "INTRODUCTION",
                    "1. First",
                    "PART 1. SETUP",
                    "2. Second",
                    "CONCLUSION",
                ],
            )
            result = ingest_script_docx(path)

            self.assertEqual(result.header_text, "INTRODUCTION")
            self.assertEqual(
                [paragraph.paragraph_no for paragraph in result.paragraphs], [1, 2]
            )
            self.assertEqual(
                [paragraph.original_index for paragraph in result.paragraphs], [2, 4]
            )
            self.assertEqual(result.issues, [])

    def test_ingest_accepts_implicit_paragraphs_and_ignores_headings(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = _write_docx(
                temp_dir,
                [
                    "INTRODUCTION",
                    "First scene",
                    "PART 1. SETUP",
                    "Second scene",
                    "CONCLUSION",
                ],
            )
            result = ingest_script_docx(path)

            self.assertEqual(result.header_text, "INTRODUCTION")
            self.assertEqual(
                [paragraph.paragraph_no for paragraph in result.paragraphs], [1, 2]
            )
            self.assertEqual(
                [paragraph.text for paragraph in result.paragraphs],
                ["First scene", "Second scene"],
            )
            self.assertEqual(result.issues, [])

    def test_keyword_extractor_fails_before_model_setup_on_invalid_numbering(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = _write_docx(temp_dir, ["3. Invalid start"])

            with self.assertRaisesRegex(ValueError, "start at 1"):
                keyword_extractor.run_keyword_extraction(path)
