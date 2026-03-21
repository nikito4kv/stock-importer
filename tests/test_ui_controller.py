from __future__ import annotations

import tempfile
import time
import unittest
from pathlib import Path
from typing import cast
from unittest.mock import patch

from docx import Document

from browser import (
    BrowserChannelResolver,
    BrowserSessionManager,
    ChromiumProfileImportService,
    ManualInterventionRequest,
    NativeBrowserSession,
    PersistentBrowserHandle,
    StoryblocksSessionProbe,
)
from domain.enums import (
    AssetKind,
    EventLevel,
    ProviderCapability,
    RunStage,
    RunStatus,
    SessionHealth,
)
from domain.models import AssetCandidate, AssetSelection, Preset, ProviderResult
from pipeline import CallbackCandidateSearchBackend
from services.errors import AppError
from services.events import AppEvent
from ui.controller import DesktopGuiController


def _write_docx(directory: str, paragraphs: list[str]) -> Path:
    path = Path(directory) / "gui-script.docx"
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(str(path))
    return path


def _candidate(
    asset_id: str, provider_name: str, kind: AssetKind, rank_hint: float
) -> AssetCandidate:
    return AssetCandidate(
        asset_id=asset_id,
        provider_name=provider_name,
        kind=kind,
        source_url=f"https://example.com/{asset_id}",
        license_name="test-license",
        metadata={"title": asset_id, "rank_hint": rank_hint},
    )


def _mark_storyblocks_ready(controller: DesktopGuiController) -> None:
    profile = (
        controller.application.container.profile_registry.get_or_create_singleton()
    )
    state = controller.application.container.session_manager.set_health(
        SessionHealth.READY
    )
    state.persistent_context_ready = True
    state.storyblocks_account = "editor@example.com"
    controller.application.container.profile_registry.update_storyblocks_account(
        profile.profile_id,
        "editor@example.com",
    )


class _FakeSessionPage:
    def __init__(self, html: str, url: str = "https://www.storyblocks.com/dashboard"):
        self._html = html
        self.url = url

    def content(self) -> str:
        return self._html

    def set_default_timeout(self, _timeout: int) -> None:
        return None

    def set_default_navigation_timeout(self, _timeout: int) -> None:
        return None

    def wait_for_load_state(self, _state: str, timeout: int | None = None) -> None:
        return None


class _FakeContextFactory:
    def __init__(self, page: _FakeSessionPage):
        self.page = page

    def launch(self, _plan):
        return PersistentBrowserHandle(
            context=object(), page=self.page, close_callback=lambda: None
        )


class _FakeNativeProcess:
    def __init__(self):
        self.running = True

    def poll(self):
        return None if self.running else 0

    def terminate(self) -> None:
        self.running = False


class _FakeNativeBrowserLauncher:
    def __init__(self):
        self.plans = []
        self.sessions: list[NativeBrowserSession] = []

    def launch(self, plan):
        self.plans.append(plan)
        session = NativeBrowserSession(plan=plan, process=_FakeNativeProcess())
        self.sessions.append(session)
        return session


class UiControllerTests(unittest.TestCase):
    def test_controller_builds_preview_and_manages_presets_and_gemini_key(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            script_path = _write_docx(
                temp_dir, ["HEADER", "1. First scene", "2. Second scene"]
            )
            controller = DesktopGuiController.create(temp_dir)

            project_summary = controller.open_script(
                script_path, project_name="GUI Project"
            )
            quick = controller.build_quick_launch_settings()
            advanced = controller.build_advanced_settings()
            quick.project_name = "GUI Project"
            quick.script_path = str(script_path)
            quick.selected_paragraphs = [1]

            preview = controller.build_run_preview(
                project_summary.project_id, quick, advanced
            )
            self.assertEqual(preview.project_name, "GUI Project")
            self.assertEqual(preview.selected_paragraphs, 1)
            self.assertIn("Ключ Gemini пока не задан", preview.warnings)

            workbench = controller.build_paragraph_workbench(project_summary.project_id)
            self.assertTrue(workbench[0].video_queries)
            self.assertTrue(workbench[0].image_queries)

            preset = controller.save_preset("quick-start", quick, advanced)
            exported = controller.export_preset(
                preset.name, Path(temp_dir) / "preset-export.json"
            )
            imported = controller.import_preset(exported)
            loaded_quick, loaded_advanced = controller.load_preset(imported.name)

            self.assertEqual(exported.name, "preset-export.json")
            self.assertEqual(imported.name, "quick-start")
            self.assertIn("quick_launch", preset.settings_snapshot)
            self.assertEqual(
                preset.settings_snapshot["launch_profile_id"],
                "normal",
            )
            self.assertNotIn("storage", preset.settings_snapshot)
            self.assertNotIn("security", preset.settings_snapshot)
            self.assertEqual(loaded_quick.mode_id, quick.mode_id)
            self.assertEqual(
                loaded_quick.launch_profile_id,
                quick.launch_profile_id,
            )
            self.assertEqual(
                loaded_advanced.action_delay_ms,
                advanced.action_delay_ms,
            )

            validation = controller.set_gemini_key("AIzaSyDUMMYKEYVALUE123456789")
            self.assertEqual(validation.severity, "success")
            self.assertEqual(
                controller.get_gemini_key(), "AIzaSyDUMMYKEYVALUE123456789"
            )

            quick.mode_id = "free_images_only"
            quick.provider_ids = ["pexels", "pixabay"]
            controller.apply_forms_to_settings(quick, advanced)

            self.assertNotIn(
                "pexels",
                controller.application.container.media_pipeline._image_backends,
            )
            pexels_saved = controller.set_provider_api_key(
                "pexels", "PEXELS_DUMMY_KEY_123456"
            )
            self.assertEqual(pexels_saved.severity, "success")
            self.assertEqual(
                controller.get_provider_api_key("pexels"), "PEXELS_DUMMY_KEY_123456"
            )
            self.assertIn(
                "pexels",
                controller.application.container.media_pipeline._image_backends,
            )

            pixabay_saved = controller.set_provider_api_key(
                "pixabay", "PIXABAY_DUMMY_KEY_123456"
            )
            self.assertEqual(pixabay_saved.severity, "success")
            self.assertEqual(
                controller.get_provider_api_key("pixabay"),
                "PIXABAY_DUMMY_KEY_123456",
            )
            self.assertIn(
                "pixabay",
                controller.application.container.media_pipeline._image_backends,
            )

            controller.delete_provider_api_key("pexels")
            self.assertIsNone(controller.get_provider_api_key("pexels"))
            self.assertNotIn(
                "pexels",
                controller.application.container.media_pipeline._image_backends,
            )

            controller.delete_provider_api_key("pixabay")
            self.assertIsNone(controller.get_provider_api_key("pixabay"))
            self.assertNotIn(
                "pixabay",
                controller.application.container.media_pipeline._image_backends,
            )

            controller.delete_gemini_key()
            self.assertIsNone(controller.get_gemini_key())

    def test_controller_projects_downloaded_files_and_run_paths(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            script_path = _write_docx(
                temp_dir,
                ["HEADER", "1. River boat scene", "2. Camp scene", "3. City scene"],
            )
            controller = DesktopGuiController.create(temp_dir)
            project = controller.open_script(script_path, project_name="Workbench")
            _mark_storyblocks_ready(controller)

            updated = controller.update_paragraph_queries(
                project.project_id,
                1,
                video_queries=["river boat dawn"],
                image_queries=["river boat photo"],
            )
            assert updated.script_document is not None
            paragraph = updated.script_document.paragraphs[0]
            assert paragraph.query_bundle is not None
            self.assertEqual(paragraph.query_bundle.video_queries, ["river boat dawn"])
            self.assertEqual(paragraph.query_bundle.image_queries, ["river boat photo"])

            container = controller.application.container
            video_descriptor = container.provider_registry.get("storyblocks_video")
            image_descriptor = container.provider_registry.get("storyblocks_image")
            assert video_descriptor is not None
            assert image_descriptor is not None

            class DownloadingVideoBackend(CallbackCandidateSearchBackend):
                def download_asset(
                    self, asset: AssetCandidate, *, destination_dir: Path, filename: str
                ) -> AssetCandidate:
                    destination_dir.mkdir(parents=True, exist_ok=True)
                    local_path = destination_dir / filename
                    local_path.write_bytes(b"video")
                    downloaded = AssetCandidate.from_dict(asset.to_dict())
                    downloaded.local_path = local_path
                    downloaded.metadata["download_status"] = "completed"
                    return downloaded

            class DownloadingImageBackend(CallbackCandidateSearchBackend):
                def download_asset(
                    self, asset: AssetCandidate, *, destination_dir: Path, filename: str
                ) -> AssetCandidate:
                    destination_dir.mkdir(parents=True, exist_ok=True)
                    local_path = destination_dir / filename
                    local_path.write_bytes(b"image")
                    downloaded = AssetCandidate.from_dict(asset.to_dict())
                    downloaded.local_path = local_path
                    downloaded.metadata["download_status"] = "completed"
                    return downloaded

            container.media_pipeline.register_backend(
                DownloadingVideoBackend(
                    provider_id="storyblocks_video",
                    capability=ProviderCapability.VIDEO,
                    descriptor=video_descriptor,
                    search_fn=lambda paragraph, query, limit: [
                        _candidate(
                            f"video-{paragraph.paragraph_no}",
                            "storyblocks_video",
                            AssetKind.VIDEO,
                            9.0,
                        )
                    ],
                )
            )
            container.media_pipeline.register_backend(
                DownloadingImageBackend(
                    provider_id="storyblocks_image",
                    capability=ProviderCapability.IMAGE,
                    descriptor=image_descriptor,
                    search_fn=lambda paragraph, query, limit: [
                        _candidate(
                            f"image-{paragraph.paragraph_no}",
                            "storyblocks_image",
                            AssetKind.IMAGE,
                            8.0,
                        )
                    ],
                )
            )

            quick = controller.build_quick_launch_settings()
            advanced = controller.build_advanced_settings()
            quick.project_name = "Workbench"
            quick.script_path = str(script_path)
            quick.mode_id = "sb_video_plus_sb_images"

            run, manifest = controller.execute_run(project.project_id, quick, advanced)
            history = controller.list_run_history(project.project_id)
            workbench = controller.build_paragraph_workbench(
                project.project_id, run.run_id
            )
            progress = controller.build_run_progress(
                run.run_id, active_project_id=project.project_id
            )

            self.assertEqual(len(history), 1)
            self.assertEqual(workbench[0].status, "completed")
            self.assertTrue(workbench[0].downloaded_files)
            self.assertTrue(
                any(
                    asset.asset_id == "video-1"
                    for asset in workbench[0].downloaded_files
                )
            )
            self.assertTrue(
                any(
                    asset.asset_id == "image-1"
                    for asset in workbench[0].downloaded_files
                )
            )
            self.assertIsNotNone(progress)
            assert progress is not None
            self.assertTrue(progress.downloads_root)
            self.assertTrue(progress.videos_dir.endswith("videos"))
            self.assertTrue(progress.images_dir.endswith("images"))
            for action_name in ("lock" + "_asset", "reject" + "_asset"):
                self.assertNotIn(action_name, dir(controller))

            rerun_run, rerun_manifest = controller.rerun_full_run(
                project.project_id, quick, advanced
            )
            self.assertEqual(
                [entry.paragraph_no for entry in rerun_manifest.paragraph_entries],
                [1, 2, 3],
            )
            self.assertEqual(rerun_run.project_id, project.project_id)

    def test_preview_warns_when_selected_free_image_providers_are_unavailable(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            script_path = _write_docx(temp_dir, ["HEADER", "1. River boat scene"])
            controller = DesktopGuiController.create(temp_dir)
            project = controller.open_script(script_path, project_name="Preview")

            quick = controller.build_quick_launch_settings()
            advanced = controller.build_advanced_settings()
            quick.mode_id = "free_images_only"
            quick.provider_ids = ["pexels", "pixabay"]

            preview = controller.build_run_preview(project.project_id, quick, advanced)

            self.assertTrue(
                any(
                    "Pexels" in warning and "API key" in warning
                    for warning in preview.warnings
                )
            )
            self.assertTrue(
                any(
                    "Pixabay" in warning and "API key" in warning
                    for warning in preview.warnings
                )
            )
            self.assertTrue(
                any(
                    "нет ни одного доступного" in warning.lower()
                    for warning in preview.warnings
                )
            )

    def test_execute_run_blocks_free_image_mode_without_usable_provider(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            script_path = _write_docx(temp_dir, ["HEADER", "1. River boat scene"])
            controller = DesktopGuiController.create(temp_dir)
            project = controller.open_script(script_path, project_name="No Providers")

            quick = controller.build_quick_launch_settings()
            advanced = controller.build_advanced_settings()
            quick.mode_id = "free_images_only"
            quick.provider_ids = ["pexels", "pixabay"]

            with self.assertRaises(AppError) as error:
                controller.execute_run(project.project_id, quick, advanced)

            self.assertEqual(error.exception.code, "free_image_provider_unavailable")

    def test_controller_supports_start_from_paragraph_range(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            script_path = _write_docx(
                temp_dir,
                ["HEADER", "1. River boat scene", "2. Camp scene", "3. City scene"],
            )
            controller = DesktopGuiController.create(temp_dir)
            project = controller.open_script(script_path, project_name="Range Start")
            _mark_storyblocks_ready(controller)

            container = controller.application.container
            video_descriptor = container.provider_registry.get("storyblocks_video")
            assert video_descriptor is not None
            container.media_pipeline.register_backend(
                CallbackCandidateSearchBackend(
                    provider_id="storyblocks_video",
                    capability=ProviderCapability.VIDEO,
                    descriptor=video_descriptor,
                    search_fn=lambda paragraph, query, limit: [
                        _candidate(
                            f"video-{paragraph.paragraph_no}",
                            "storyblocks_video",
                            AssetKind.VIDEO,
                            9.0,
                        )
                    ],
                )
            )

            quick = controller.build_quick_launch_settings()
            advanced = controller.build_advanced_settings()
            quick.mode_id = "sb_video_only"
            quick.paragraph_selection_text = "2..end"

            preview = controller.build_run_preview(project.project_id, quick, advanced)
            run, manifest = controller.execute_run(project.project_id, quick, advanced)

            self.assertEqual(preview.selected_paragraphs, 2)
            self.assertEqual(run.selected_paragraphs, [2, 3])
            self.assertEqual(
                [entry.paragraph_no for entry in manifest.paragraph_entries], [2, 3]
            )

    def test_controller_manages_storyblocks_session_panel_without_cli(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            controller = DesktopGuiController.create(temp_dir)
            profile = controller.application.container.profile_registry.get_or_create_singleton()
            container = controller.application.container

            def fake_open_browser():
                state = container.session_manager.set_health(SessionHealth.UNKNOWN)
                state.persistent_context_ready = True
                return cast(object, None)

            def fake_native_login(**_kwargs):
                state = container.session_manager.set_health(
                    SessionHealth.LOGIN_REQUIRED
                )
                state.manual_intervention = ManualInterventionRequest(
                    reason="native_login",
                    prompt="Close the browser window used for Storyblocks login, then click Check Session again.",
                    requested_at=state.last_checked_at,
                )
                return state

            def fake_check_authorization(**_kwargs):
                state = container.session_manager.set_health(SessionHealth.READY)
                state.storyblocks_account = "editor@example.com"
                container.profile_registry.update_storyblocks_account(
                    profile.profile_id,
                    "editor@example.com",
                )
                return state

            container.session_manager.open_native_login_browser = fake_native_login
            container.session_manager.check_authorization = fake_check_authorization

            initial = controller.session_panel()
            login = controller.prepare_storyblocks_login()
            ready = controller.check_storyblocks_session()

            self.assertEqual(initial.health, SessionHealth.UNKNOWN.value)
            self.assertEqual(login.health, SessionHealth.LOGIN_REQUIRED.value)
            self.assertEqual(ready.health, SessionHealth.READY.value)
            self.assertEqual(ready.account, "editor@example.com")

    def test_controller_can_reset_storyblocks_session_with_reason_code(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            controller = DesktopGuiController.create(temp_dir)
            _mark_storyblocks_ready(controller)

            panel = controller.reset_storyblocks_session()

            self.assertEqual(panel.health, SessionHealth.LOGIN_REQUIRED.value)
            self.assertEqual(panel.reason_code, "session_reset")
            self.assertIn("Log in again", panel.last_error)

    def test_controller_allows_manual_storyblocks_override_for_run_gating(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            script_path = _write_docx(temp_dir, ["HEADER", "1. River boat scene"])
            controller = DesktopGuiController.create(temp_dir)
            project = controller.open_script(script_path, project_name="Override Run")
            controller.application.container.session_manager.set_health(
                SessionHealth.UNKNOWN
            )
            panel = controller.mark_storyblocks_session_ready()

            self.assertTrue(panel.manual_ready_override)

            container = controller.application.container
            video_descriptor = container.provider_registry.get("storyblocks_video")
            assert video_descriptor is not None
            container.media_pipeline.register_backend(
                CallbackCandidateSearchBackend(
                    provider_id="storyblocks_video",
                    capability=ProviderCapability.VIDEO,
                    descriptor=video_descriptor,
                    search_fn=lambda paragraph, query, limit: [
                        _candidate(
                            f"video-{paragraph.paragraph_no}",
                            "storyblocks_video",
                            AssetKind.VIDEO,
                            10.0,
                        )
                    ],
                )
            )

            quick = controller.build_quick_launch_settings()
            advanced = controller.build_advanced_settings()
            quick.project_name = "Override Run"
            quick.script_path = str(script_path)
            quick.mode_id = "sb_video_only"

            run, _manifest = controller.execute_run(project.project_id, quick, advanced)

            self.assertEqual(run.status.value, "completed")
            cleared = controller.clear_storyblocks_session_override()
            self.assertFalse(cleared.manual_ready_override)

    def test_controller_can_import_existing_chrome_session(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            external_root = (
                Path(temp_dir) / "external" / "Google" / "Chrome" / "User Data"
            )
            source_profile = external_root / "Default"
            (source_profile / "Network").mkdir(parents=True, exist_ok=True)
            (source_profile / "Preferences").write_text(
                '{"profile": {"name": "Storyblocks Personal"}}', encoding="utf-8"
            )
            (source_profile / "Network" / "Cookies").write_bytes(b"cookies")
            (external_root / "Local State").write_text(
                '{"os_crypt": {"encrypted_key": "dummy"}}', encoding="utf-8"
            )

            controller = DesktopGuiController.create(temp_dir)
            container = controller.application.container
            container.profile_import_service = ChromiumProfileImportService(
                container.profile_registry,
                explicit_user_data_roots={"chrome": [external_root]},
            )
            chrome_path = Path(temp_dir) / "chrome.exe"
            chrome_path.write_bytes(b"x")
            container.session_manager = BrowserSessionManager(
                container.profile_registry,
                container.settings.browser,
                channel_resolver=BrowserChannelResolver(
                    explicit_candidates={"chrome": [chrome_path]}
                ),
                context_factory=_FakeContextFactory(
                    _FakeSessionPage(
                        html='<div data-account-email="editor@example.com"></div><button>Download</button><a href="/logout">Logout</a>'
                    )
                ),
                session_probe=StoryblocksSessionProbe(),
            )

            options = controller.discover_importable_browser_profiles("chrome")
            panel = controller.import_storyblocks_session_from_path(
                options[0].profile_dir, browser_name="chrome"
            )

            self.assertEqual(panel.health, SessionHealth.READY.value)
            self.assertIn(str(external_root), panel.imported_source)
            self.assertEqual(panel.imported_profile_name, "Storyblocks Personal")

    def test_controller_blocks_run_when_numbering_is_invalid(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            script_path = _write_docx(
                temp_dir, ["HEADER", "2. Broken numbering", "3. Another paragraph"]
            )
            controller = DesktopGuiController.create(temp_dir)
            project = controller.open_script(
                script_path, project_name="Broken Numbering"
            )
            _mark_storyblocks_ready(controller)

            with self.assertRaises(AppError) as ctx:
                controller.start_run_async(
                    project.project_id,
                    controller.build_quick_launch_settings(),
                    controller.build_advanced_settings(),
                )

            self.assertEqual(ctx.exception.code, "invalid_numbering")

    def test_controller_blocks_storyblocks_parallelism_over_one_worker(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            script_path = _write_docx(temp_dir, ["HEADER", "1. River boat scene"])
            controller = DesktopGuiController.create(temp_dir)
            controller.open_script(script_path, project_name="Guarded Run")
            _mark_storyblocks_ready(controller)

            quick = controller.build_quick_launch_settings()
            advanced = controller.build_advanced_settings()
            quick.mode_id = "sb_video_only"
            quick.launch_profile_id = "fast"

            settings = controller.apply_forms_to_settings(quick, advanced)

            self.assertEqual(settings.concurrency.paragraph_workers, 1)
            self.assertEqual(settings.concurrency.queue_size, 1)
            self.assertEqual(
                controller.application.container.orchestrator.max_workers,
                1,
            )

    def test_controller_allows_free_image_parallelism_over_one_worker(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            script_path = _write_docx(temp_dir, ["HEADER", "1. River boat scene"])
            controller = DesktopGuiController.create(temp_dir)
            project = controller.open_script(
                script_path, project_name="Free Images Run"
            )

            quick = controller.build_quick_launch_settings()
            advanced = controller.build_advanced_settings()
            quick.mode_id = "free_images_only"
            quick.launch_profile_id = "fast"
            quick.provider_ids = ["openverse"]

            with patch.object(controller, "_start_background_run", return_value=None):
                run_id = controller.start_run_async(project.project_id, quick, advanced)

            self.assertTrue(run_id)
            self.assertIsNotNone(
                controller.application.container.run_repository.load(run_id)
            )
            self.assertEqual(
                controller.application.container.orchestrator.max_workers,
                4,
            )

    def test_build_quick_launch_settings_infers_normal_fast_and_custom_profiles(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            controller = DesktopGuiController.create(temp_dir)

            self.assertEqual(
                controller.build_quick_launch_settings().launch_profile_id,
                "normal",
            )

            quick = controller.build_quick_launch_settings()
            advanced = controller.build_advanced_settings()
            quick.launch_profile_id = "fast"
            controller.apply_forms_to_settings(quick, advanced)

            self.assertEqual(
                controller.build_quick_launch_settings().launch_profile_id,
                "fast",
            )

            quick = controller.build_quick_launch_settings()
            advanced = controller.build_advanced_settings()
            quick.launch_profile_id = "custom"
            advanced.launch_timeout_ms = 61000
            advanced.navigation_timeout_ms = 47000
            advanced.downloads_timeout_seconds = 240.0
            controller.apply_forms_to_settings(quick, advanced)

            inferred_quick = controller.build_quick_launch_settings()
            inferred_advanced = controller.build_advanced_settings()
            self.assertEqual(inferred_quick.launch_profile_id, "custom")
            self.assertEqual(inferred_advanced.launch_timeout_ms, 61000)
            self.assertEqual(inferred_advanced.navigation_timeout_ms, 47000)
            self.assertEqual(inferred_advanced.downloads_timeout_seconds, 240.0)

    def test_media_config_from_forms_uses_resolved_launch_profile_timeouts(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            controller = DesktopGuiController.create(temp_dir)
            quick = controller.build_quick_launch_settings()
            advanced = controller.build_advanced_settings()

            normal_config = controller._media_config_from_forms(quick, advanced)
            quick.launch_profile_id = "fast"
            fast_config = controller._media_config_from_forms(quick, advanced)

            self.assertEqual(normal_config.search_timeout_seconds, 20.0)
            self.assertEqual(fast_config.search_timeout_seconds, 12.0)
            self.assertEqual(normal_config.download_timeout_seconds, 120.0)
            self.assertEqual(fast_config.download_timeout_seconds, 90.0)
            self.assertEqual(normal_config.retry_budget, 2)
            self.assertEqual(fast_config.retry_budget, 1)

    def test_load_legacy_preset_maps_to_custom_timing_and_safe_runtime_defaults(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            controller = DesktopGuiController.create(temp_dir)
            legacy = Preset(
                name="legacy-custom",
                settings_snapshot={
                    "browser": {
                        "slow_mode": True,
                        "action_delay_ms": 1500,
                        "launch_timeout_ms": 55000,
                        "navigation_timeout_ms": 36000,
                        "downloads_timeout_seconds": 180.0,
                    },
                    "providers": {
                        "project_mode": "sb_video_only",
                    },
                    "concurrency": {
                        "paragraph_workers": 8,
                        "queue_size": 8,
                        "search_timeout_seconds": 99.0,
                        "relevance_timeout_seconds": 33.0,
                        "retry_budget": 9,
                    },
                },
            )
            controller.application.container.settings_manager.save_preset(legacy)

            quick, advanced = controller.load_preset(legacy.name)

            self.assertEqual(quick.launch_profile_id, "custom")
            self.assertEqual(advanced.action_delay_ms, 1500)
            self.assertEqual(advanced.launch_timeout_ms, 55000)
            self.assertEqual(advanced.navigation_timeout_ms, 36000)
            self.assertEqual(advanced.downloads_timeout_seconds, 180.0)

            settings = controller.application.container.settings
            self.assertEqual(settings.concurrency.paragraph_workers, 1)
            self.assertEqual(settings.concurrency.queue_size, 1)
            self.assertEqual(settings.browser.action_delay_ms, 1500)
            self.assertTrue(settings.browser.slow_mode)
            self.assertFalse(hasattr(settings.concurrency, "relevance_timeout_seconds"))

    def test_custom_preset_round_trip_preserves_launch_profile_and_timing(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            controller = DesktopGuiController.create(temp_dir)
            quick = controller.build_quick_launch_settings()
            advanced = controller.build_advanced_settings()
            quick.launch_profile_id = "custom"
            advanced.action_delay_ms = 250
            advanced.launch_timeout_ms = 60000
            advanced.navigation_timeout_ms = 45000
            advanced.downloads_timeout_seconds = 240.0

            preset = controller.save_preset("custom-round-trip", quick, advanced)
            exported = controller.export_preset(
                preset.name, Path(temp_dir) / "custom-round-trip.json"
            )
            imported = controller.import_preset(exported)
            loaded_quick, loaded_advanced = controller.load_preset(imported.name)

            self.assertEqual(preset.settings_snapshot["launch_profile_id"], "custom")
            self.assertEqual(
                preset.settings_snapshot["custom_timing_overrides"],
                {
                    "action_delay_ms": 250,
                    "launch_timeout_ms": 60000,
                    "navigation_timeout_ms": 45000,
                    "downloads_timeout_seconds": 240.0,
                },
            )
            self.assertNotIn("storage", preset.settings_snapshot)
            self.assertNotIn("security", preset.settings_snapshot)
            self.assertNotIn("cache_root", preset.settings_snapshot["quick_launch"])
            self.assertNotIn(
                "browser_profile_path", preset.settings_snapshot["quick_launch"]
            )
            self.assertEqual(loaded_quick.launch_profile_id, "custom")
            self.assertEqual(loaded_advanced.action_delay_ms, 250)
            self.assertEqual(loaded_advanced.launch_timeout_ms, 60000)
            self.assertEqual(loaded_advanced.navigation_timeout_ms, 45000)
            self.assertEqual(loaded_advanced.downloads_timeout_seconds, 240.0)

    def test_build_live_snapshot_does_not_require_manifest_reload(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            script_path = _write_docx(temp_dir, ["HEADER", "1. River boat scene"])
            controller = DesktopGuiController.create(temp_dir)
            project = controller.open_script(script_path, project_name="Live Snapshot")
            _mark_storyblocks_ready(controller)
            container = controller.application.container
            descriptor = container.provider_registry.get("storyblocks_video")
            assert descriptor is not None

            def slow_video(paragraph, query, limit):
                time.sleep(0.2)
                return [
                    _candidate(
                        f"video-{paragraph.paragraph_no}",
                        "storyblocks_video",
                        AssetKind.VIDEO,
                        10.0,
                    )
                ]

            container.media_pipeline.register_backend(
                CallbackCandidateSearchBackend(
                    provider_id="storyblocks_video",
                    capability=ProviderCapability.VIDEO,
                    descriptor=descriptor,
                    search_fn=slow_video,
                )
            )

            quick = controller.build_quick_launch_settings()
            advanced = controller.build_advanced_settings()
            quick.mode_id = "sb_video_only"
            run_id = controller.start_run_async(project.project_id, quick, advanced)

            with (
                patch.object(
                    controller.application.container.media_run_service,
                    "snapshot_manifest",
                    side_effect=AssertionError("manifest snapshot should not happen"),
                ),
                patch.object(
                    controller.application.container.media_run_service,
                    "load_manifest",
                    side_effect=AssertionError("manifest load should not happen"),
                ),
                patch.object(
                    controller.application.container.project_repository,
                    "load",
                    side_effect=AssertionError("project load should not happen"),
                ),
            ):
                snapshot = controller.build_live_snapshot(
                    active_project_id=project.project_id,
                    active_run_id=run_id,
                )
                self.assertEqual(snapshot.active_run_id, run_id)
                self.assertIsNotNone(snapshot.run_progress)

            deadline = time.time() + 10.0
            while time.time() < deadline:
                state = controller.build_state(
                    active_project_id=project.project_id, active_run_id=run_id
                )
                if state.run_progress is not None and state.run_progress.status in {
                    "completed",
                    "failed",
                    "cancelled",
                }:
                    break
                time.sleep(0.05)

    def test_build_live_run_state_uses_compact_snapshot_and_builds_workbench_once(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            script_path = _write_docx(
                temp_dir, ["HEADER", "1. River boat scene", "2. Camp scene"]
            )
            controller = DesktopGuiController.create(temp_dir)
            project = controller.open_script(script_path, project_name="Live Run State")
            container = controller.application.container
            run, manifest = container.media_run_service.create_run(project.project_id)
            run.status = RunStatus.RUNNING
            run.stage = RunStage.PROVIDER_SEARCH
            container.run_repository.save(run)

            detail_candidate = _candidate(
                "video-1", "storyblocks_video", AssetKind.VIDEO, 10.0
            )
            detail_candidate.local_path = Path(temp_dir) / "downloads" / "video-1.mp4"
            detail_candidate.local_path.parent.mkdir(parents=True, exist_ok=True)
            detail_candidate.local_path.write_bytes(b"video")
            detail_entry = manifest.paragraph_entries[0]
            detail_entry.selection = AssetSelection(
                paragraph_no=1,
                primary_asset=detail_candidate,
                reason="Saved video is ready",
                provider_results=[
                    ProviderResult(
                        provider_name="storyblocks_video",
                        capability=ProviderCapability.VIDEO,
                        query="river boat",
                        candidates=[detail_candidate],
                    )
                ],
                status="completed",
            )
            detail_entry.status = "completed"
            container.media_pipeline.register_live_manifest(manifest)
            container.event_recorder(
                AppEvent(
                    name="provider.search.started",
                    level=EventLevel.INFO,
                    message="Provider search started",
                    project_id=project.project_id,
                    run_id=run.run_id,
                    paragraph_no=1,
                    provider_name="storyblocks_video",
                    query="river boat",
                    stage=RunStage.PROVIDER_SEARCH,
                )
            )

            with (
                patch.object(
                    controller.application.container.media_run_service,
                    "snapshot_live_run_state",
                    wraps=(
                        controller.application.container.media_run_service.snapshot_live_run_state
                    ),
                ) as snapshot_live_run_state,
                patch.object(
                    controller.application.container.media_run_service,
                    "snapshot_manifest",
                    side_effect=AssertionError("manifest snapshot should not happen"),
                ),
                patch.object(
                    controller.application.container.media_run_service,
                    "load_manifest",
                    side_effect=AssertionError("manifest load should not happen"),
                ),
                patch.object(
                    controller.application.container.project_repository,
                    "load",
                    side_effect=AssertionError("project load should not happen"),
                ),
                patch.object(
                    controller,
                    "build_paragraph_workbench",
                    wraps=controller.build_paragraph_workbench,
                ) as build_paragraph_workbench,
            ):
                live_state = controller.build_live_run_state(
                    active_project_id=project.project_id,
                    active_run_id=run.run_id,
                )

            snapshot_live_run_state.assert_called_once_with(
                run.run_id, detailed_paragraph_no=1
            )
            build_paragraph_workbench.assert_called_once()
            self.assertIsNotNone(live_state.run_progress)
            self.assertEqual(len(live_state.paragraph_items), 2)
            item_by_no = {
                item.paragraph_no: item for item in live_state.paragraph_items
            }
            self.assertTrue(item_by_no[1].text)
            self.assertEqual(item_by_no[2].text, "")
            self.assertEqual(item_by_no[1].status, "completed")
            self.assertEqual(item_by_no[1].result_note, "Saved video is ready")
            self.assertEqual(len(item_by_no[1].downloaded_files), 1)
            self.assertEqual(
                item_by_no[1].downloaded_files[0].local_path,
                str(detail_candidate.local_path),
            )
            for field_name in ("selected" + "_assets", "candidate" + "_assets"):
                self.assertNotIn(field_name, type(item_by_no[1]).__slots__)

    def test_build_live_run_state_does_not_fallback_to_manifest_after_live_release(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            script_path = _write_docx(temp_dir, ["HEADER", "1. River boat scene"])
            controller = DesktopGuiController.create(temp_dir)
            project = controller.open_script(script_path, project_name="Released Live")
            container = controller.application.container
            run, manifest = container.media_run_service.create_run(project.project_id)
            run.status = RunStatus.COMPLETED
            run.stage = RunStage.COMPLETE
            container.run_repository.save(run)
            container.media_pipeline.register_live_manifest(manifest)
            container.media_pipeline.release_run_state(run.run_id)
            container.event_recorder(
                AppEvent(
                    name="run.completed",
                    level=EventLevel.INFO,
                    message="Run completed",
                    project_id=project.project_id,
                    run_id=run.run_id,
                    stage=RunStage.COMPLETE,
                )
            )

            with (
                patch.object(
                    controller.application.container.media_run_service,
                    "snapshot_live_run_state",
                    wraps=(
                        controller.application.container.media_run_service.snapshot_live_run_state
                    ),
                ) as snapshot_live_run_state,
                patch.object(
                    controller.application.container.media_run_service,
                    "snapshot_manifest",
                    side_effect=AssertionError("manifest snapshot should not happen"),
                ),
                patch.object(
                    controller.application.container.media_run_service,
                    "load_manifest",
                    side_effect=AssertionError("manifest load should not happen"),
                ),
            ):
                live_state = controller.build_live_run_state(
                    active_project_id=project.project_id,
                    active_run_id=run.run_id,
                )

            snapshot_live_run_state.assert_called_once_with(
                run.run_id, detailed_paragraph_no=None
            )
            self.assertIsNotNone(live_state.run_progress)
            assert live_state.run_progress is not None
            self.assertEqual(live_state.run_progress.status, "completed")
            self.assertEqual(live_state.paragraph_items, [])

    def test_controller_blocks_session_changes_while_run_is_active(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            script_path = _write_docx(temp_dir, ["HEADER", "1. River boat scene"])
            controller = DesktopGuiController.create(temp_dir)
            project = controller.open_script(script_path, project_name="Busy Run")
            _mark_storyblocks_ready(controller)
            container = controller.application.container
            video_descriptor = container.provider_registry.get("storyblocks_video")
            assert video_descriptor is not None

            def slow_video(paragraph, query, limit):
                time.sleep(0.35)
                return [
                    _candidate(
                        f"video-{paragraph.paragraph_no}",
                        "storyblocks_video",
                        AssetKind.VIDEO,
                        10.0,
                    )
                ]

            container.media_pipeline.register_backend(
                CallbackCandidateSearchBackend(
                    provider_id="storyblocks_video",
                    capability=ProviderCapability.VIDEO,
                    descriptor=video_descriptor,
                    search_fn=slow_video,
                )
            )

            quick = controller.build_quick_launch_settings()
            advanced = controller.build_advanced_settings()
            quick.project_name = "Busy Run"
            quick.script_path = str(script_path)
            quick.mode_id = "sb_video_only"

            run_id = controller.start_run_async(project.project_id, quick, advanced)
            time.sleep(0.05)

            with self.assertRaises(AppError) as ctx:
                controller.prepare_storyblocks_login()

            self.assertEqual(ctx.exception.code, "run_in_progress")

            deadline = time.time() + 10.0
            while time.time() < deadline:
                state = controller.build_state(
                    active_project_id=project.project_id, active_run_id=run_id
                )
                if state.run_progress is not None and state.run_progress.status in {
                    "completed",
                    "failed",
                    "cancelled",
                }:
                    break
                time.sleep(0.05)

    def test_controller_tracks_live_progress_without_event_journal(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            script_path = _write_docx(
                temp_dir, ["HEADER", "1. River boat scene", "2. Camp scene"]
            )
            controller = DesktopGuiController.create(temp_dir)
            project = controller.open_script(script_path, project_name="Live Progress")
            _mark_storyblocks_ready(controller)
            container = controller.application.container
            video_descriptor = container.provider_registry.get("storyblocks_video")
            image_descriptor = container.provider_registry.get("storyblocks_image")
            assert video_descriptor is not None
            assert image_descriptor is not None

            def slow_video(paragraph, query, limit):
                time.sleep(0.35)
                return [
                    _candidate(
                        f"video-{paragraph.paragraph_no}",
                        "storyblocks_video",
                        AssetKind.VIDEO,
                        10.0,
                    )
                ]

            def slow_image(paragraph, query, limit):
                time.sleep(0.05)
                return [
                    _candidate(
                        f"image-{paragraph.paragraph_no}",
                        "storyblocks_image",
                        AssetKind.IMAGE,
                        8.0,
                    )
                ]

            container.media_pipeline.register_backend(
                CallbackCandidateSearchBackend(
                    provider_id="storyblocks_video",
                    capability=ProviderCapability.VIDEO,
                    descriptor=video_descriptor,
                    search_fn=slow_video,
                )
            )
            container.media_pipeline.register_backend(
                CallbackCandidateSearchBackend(
                    provider_id="storyblocks_image",
                    capability=ProviderCapability.IMAGE,
                    descriptor=image_descriptor,
                    search_fn=slow_image,
                )
            )

            quick = controller.build_quick_launch_settings()
            advanced = controller.build_advanced_settings()
            quick.project_name = "Live Progress"
            quick.script_path = str(script_path)
            quick.mode_id = "sb_video_plus_sb_images"

            run_id = controller.start_run_async(project.project_id, quick, advanced)

            live_snapshot = None
            deadline = time.time() + 10.0
            while time.time() < deadline:
                live_snapshot = controller.build_live_snapshot(
                    active_project_id=project.project_id, active_run_id=run_id
                )
                if (
                    live_snapshot.run_progress is not None
                    and live_snapshot.run_progress.status == "running"
                ):
                    break
                time.sleep(0.05)

            assert live_snapshot is not None
            assert live_snapshot.run_progress is not None
            self.assertEqual(live_snapshot.run_progress.status, "running")
            self.assertTrue(live_snapshot.run_progress.can_cancel)
            self.assertFalse(hasattr(live_snapshot, "event_journal"))

            live_state = controller.build_live_run_state(
                active_project_id=project.project_id,
                active_run_id=run_id,
                live_snapshot=live_snapshot,
            )
            self.assertFalse(hasattr(live_state, "event_journal"))

            running_state = controller.build_state(
                active_project_id=project.project_id,
                active_run_id=run_id,
            )
            self.assertFalse(hasattr(running_state, "event_journal"))

            controller.cancel_run(run_id)

            completed_state = None
            deadline = time.time() + 20.0
            while time.time() < deadline:
                completed_state = controller.build_state(
                    active_project_id=project.project_id, active_run_id=run_id
                )
                if (
                    completed_state.run_progress is not None
                    and completed_state.run_progress.status == "cancelled"
                ):
                    break
                time.sleep(0.05)

            assert completed_state is not None
            assert completed_state.run_progress is not None
            self.assertEqual(completed_state.run_progress.status, "cancelled")
            history = controller.list_run_history(project.project_id)
            self.assertEqual(history[0].status, "cancelled")

    def test_controller_rerun_full_run_async_ignores_current_selection(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            script_path = _write_docx(
                temp_dir,
                ["HEADER", "1. River boat scene", "2. Camp scene", "3. City scene"],
            )
            controller = DesktopGuiController.create(temp_dir)
            project = controller.open_script(script_path, project_name="Full Rerun")
            _mark_storyblocks_ready(controller)
            container = controller.application.container
            video_descriptor = container.provider_registry.get("storyblocks_video")
            assert video_descriptor is not None

            container.media_pipeline.register_backend(
                CallbackCandidateSearchBackend(
                    provider_id="storyblocks_video",
                    capability=ProviderCapability.VIDEO,
                    descriptor=video_descriptor,
                    search_fn=lambda paragraph, query, limit: [
                        _candidate(
                            f"video-{paragraph.paragraph_no}",
                            "storyblocks_video",
                            AssetKind.VIDEO,
                            9.0,
                        )
                    ],
                )
            )

            quick = controller.build_quick_launch_settings()
            advanced = controller.build_advanced_settings()
            quick.project_name = "Full Rerun"
            quick.script_path = str(script_path)
            quick.mode_id = "sb_video_only"
            quick.paragraph_selection_text = "2..end"

            subset_run, subset_manifest = controller.execute_run(
                project.project_id,
                quick,
                advanced,
            )
            self.assertEqual(subset_run.selected_paragraphs, [2, 3])
            self.assertEqual(
                [entry.paragraph_no for entry in subset_manifest.paragraph_entries],
                [2, 3],
            )

            quick.paragraph_selection_text = "1"
            rerun_id = controller.rerun_full_run_async(
                project.project_id,
                quick,
                advanced,
            )

            rerun_state = None
            deadline = time.time() + 20.0
            while time.time() < deadline:
                rerun_state = controller.build_state(
                    active_project_id=project.project_id,
                    active_run_id=rerun_id,
                )
                if (
                    rerun_state.run_progress is not None
                    and rerun_state.run_progress.status == "completed"
                ):
                    break
                time.sleep(0.05)

            rerun_run = controller.application.container.run_repository.load(rerun_id)
            rerun_manifest = controller.application.container.manifest_repository.load(
                rerun_id
            )
            self.assertIsNotNone(rerun_run)
            self.assertIsNotNone(rerun_manifest)
            assert rerun_run is not None
            assert rerun_manifest is not None
            self.assertEqual(rerun_run.selected_paragraphs, [])
            self.assertEqual(
                [entry.paragraph_no for entry in rerun_manifest.paragraph_entries],
                [1, 2, 3],
            )


if __name__ == "__main__":
    unittest.main()
